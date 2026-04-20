"""
DOCX importer.

Uses python-docx to extract body text from .docx files. Paragraphs
and tables are walked in document order (not "all paragraphs then all
tables") so the reading flow matches the original layout.

Tables are flattened row-by-row: each row becomes a line with cell
texts separated by tabs. This reads reasonably in RSVP even though
it loses column alignment.

Headers, footers, footnotes, and endnotes are skipped — they're
peripheral content that would interrupt the reading flow.

Requires: python-docx (listed in requirements.txt).
"""

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

from core.tokenizer import Token, tokenize
from importers.base import Importer


def _extract_text(doc: Document) -> str:
    """Walk the document body in element order, extracting text from
    paragraphs and tables as they appear.

    python-docx exposes Document.paragraphs and Document.tables as
    separate flat lists, which loses interleaving. To preserve reading
    order we iterate the underlying XML children of the body element
    and dispatch on tag name.
    """
    parts: list[str] = []
    body = doc.element.body

    for child in body.iterchildren():
        tag = child.tag

        if tag == qn("w:p"):
            # It's a paragraph element. Wrap it in the python-docx
            # Paragraph class to get its .text property.
            para = Paragraph(child, doc)
            text = para.text.strip()
            if text:
                parts.append(text)

        elif tag == qn("w:tbl"):
            # It's a table element. Flatten row-by-row.
            table = Table(child, doc)
            table_lines: list[str] = []
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                line = "\t".join(cells)
                if line.strip():
                    table_lines.append(line)
            if table_lines:
                parts.append("\n".join(table_lines))

        # Other elements (w:sectPr, w:bookmarkStart, etc.) are
        # structural/formatting and carry no readable text — skip.

    return "\n\n".join(parts)


class DocxImporter(Importer):

    @property
    def extensions(self) -> tuple[str, ...]:
        return (".docx",)

    def can_handle(self, path: Path) -> bool:
        return path.suffix.lower() in self.extensions

    def load(self, path: Path) -> tuple[str, list[Token]]:
        doc = Document(str(path))
        canonical = _extract_text(doc)
        return canonical, tokenize(canonical)
